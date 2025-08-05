import io
import re
import email
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import requests
import PyPDF2
from docx import Document
import asyncio
import aiofiles
from .config import settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles extraction and preprocessing of various document formats"""
    
    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
    
    async def process_document_from_url(self, url: str) -> Dict[str, Any]:
        """Download and process document from URL"""
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            content_type = response.headers.get('content-type', '')
            file_extension = self._get_file_extension_from_url(url, content_type)
            
            document_content = await self._extract_content(
                response.content, file_extension
            )
            
            chunks = self._create_chunks(document_content)
            
            return {
                'content': document_content,
                'chunks': chunks,
                'metadata': {
                    'source_url': url,
                    'content_type': content_type,
                    'file_extension': file_extension,
                    'total_chunks': len(chunks),
                    'content_length': len(document_content)
                }
            }
        except Exception as e:
            logger.error(f"Error processing document from URL {url}: {str(e)}")
            raise
    
    def _get_file_extension_from_url(self, url: str, content_type: str) -> str:
        """Determine file extension from URL or content type"""
        if '.pdf' in url.lower() or 'pdf' in content_type:
            return '.pdf'
        elif '.docx' in url.lower() or 'document' in content_type:
            return '.docx'
        elif '.txt' in url.lower() or 'text' in content_type:
            return '.txt'
        else:
            return '.pdf'  # default assumption
    
    async def _extract_content(self, file_content: bytes, file_extension: str) -> str:
        """Extract text content based on file type"""
        try:
            if file_extension == '.pdf':
                return await self._extract_pdf_content(file_content)
            elif file_extension == '.docx':
                return await self._extract_docx_content(file_content)
            elif file_extension == '.txt':
                return file_content.decode('utf-8', errors='ignore')
            elif file_extension == '.eml':
                return await self._extract_email_content(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
        except Exception as e:
            logger.error(f"Error extracting content from {file_extension}: {str(e)}")
            raise
    
    async def _extract_pdf_content(self, file_content: bytes) -> str:
        """Extract text from PDF content"""
        text_content = []
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"[Page {page_num + 1}]\n{page_text}")
        except Exception as e:
            logger.error(f"Error extracting PDF content: {str(e)}")
            raise
        
        return "\n\n".join(text_content)
    
    async def _extract_docx_content(self, file_content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            doc = Document(io.BytesIO(file_content))
            text_content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {str(e)}")
            raise
    
    async def _extract_email_content(self, file_content: bytes) -> str:
        """Extract text from email content"""
        try:
            msg = email.message_from_bytes(file_content)
            
            content_parts = []
            
            # Extract headers
            headers = ['From', 'To', 'Subject', 'Date']
            for header in headers:
                if msg.get(header):
                    content_parts.append(f"{header}: {msg.get(header)}")
            
            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload:
                            content_parts.append(payload.decode('utf-8', errors='ignore'))
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    content_parts.append(payload.decode('utf-8', errors='ignore'))
            
            return "\n\n".join(content_parts)
        except Exception as e:
            logger.error(f"Error extracting email content: {str(e)}")
            raise
    
    def _create_chunks(self, content: str) -> List[Dict[str, Any]]:
        """Split content into overlapping chunks"""
        sentences = self._split_into_sentences(content)
        chunks = []
        current_chunk = ""
        current_size = 0
        chunk_id = 0
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            if current_size + sentence_size > self.chunk_size and current_chunk:
                # Save current chunk
                chunks.append({
                    'chunk_id': f"chunk_{chunk_id}",
                    'content': current_chunk.strip(),
                    'size': current_size,
                    'metadata': {
                        'chunk_index': chunk_id,
                        'start_sentence': len(chunks) * (self.chunk_size // 100),  # Approximate
                    }
                })
                
                chunk_id += 1
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + " " + sentence
                current_size = len(overlap_text) + sentence_size
            else:
                current_chunk += " " + sentence
                current_size += sentence_size
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append({
                'chunk_id': f"chunk_{chunk_id}",
                'content': current_chunk.strip(),
                'size': current_size,
                'metadata': {
                    'chunk_index': chunk_id,
                    'start_sentence': len(chunks) * (self.chunk_size // 100),
                }
            })
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using regex"""
        # Simple sentence splitting - can be enhanced with NLTK
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text from the end of current chunk"""
        words = text.split()
        overlap_words = words[-self.chunk_overlap//10:]  # Approximate word count
        return " ".join(overlap_words)
    
    def preprocess_text(self, text: str) -> str:
        """Clean and preprocess text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s.,!?;:()\-]', '', text)
        
        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        return text.strip()

